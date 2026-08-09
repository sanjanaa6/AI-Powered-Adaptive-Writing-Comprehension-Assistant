import os
import pypdf
import docx

def read_text_from_file(file_path: str) -> str:
    """Reads raw text from PDF, DOCX, or TXT file."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
            
    elif ext == ".pdf":
        text = ""
        reader = pypdf.PdfReader(file_path)
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text
        
    elif ext in [".docx", ".doc"]:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text])
        
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def extract_bytes_to_text(file_bytes: bytes, file_name: str) -> str:
    """Reads raw text from uploaded bytes stream."""
    ext = os.path.splitext(file_name)[1].lower()
    
    if ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")
        
    elif ext == ".pdf":
        import io
        text = ""
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text
        
    elif ext in [".docx", ".doc"]:
        import io
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join([p.text for p in doc.paragraphs if p.text])
        
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def split_text_into_chunks(text: str, chunk_size: int = 800, chunk_overlap: int = 150) -> list:
    """Pure Python text splitter that breaks text into chunks with overlap without external dependencies."""
    if not text.strip():
        return []

    # First split by paragraphs
    paragraphs = text.split("\n\n")
    chunks = []
    current_chunk = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        if len(current_chunk) + len(para) + 2 <= chunk_size:
            current_chunk = f"{current_chunk}\n\n{para}" if current_chunk else para
        else:
            if current_chunk:
                chunks.append(current_chunk)
            
            # If paragraph itself is larger than chunk_size, split by sentences or spaces
            if len(para) > chunk_size:
                words = para.split(" ")
                sub_chunk = ""
                for w in words:
                    if len(sub_chunk) + len(w) + 1 <= chunk_size:
                        sub_chunk = f"{sub_chunk} {w}" if sub_chunk else w
                    else:
                        chunks.append(sub_chunk)
                        # Overlap: keep last few words
                        overlap_words = sub_chunk.split(" ")[-10:]
                        sub_chunk = f"{' '.join(overlap_words)} {w}"
                if sub_chunk:
                    current_chunk = sub_chunk
                else:
                    current_chunk = ""
            else:
                current_chunk = para

    if current_chunk:
        chunks.append(current_chunk)

    return chunks
